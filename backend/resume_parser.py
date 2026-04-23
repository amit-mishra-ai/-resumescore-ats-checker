"""
resume_parser.py
Extracts plain text from PDF and DOCX resume files.
Keeps it lightweight — no heavy NLP libraries needed here.
"""

import re


def extract_text_from_pdf(filepath):
    """Extract text from a PDF file using PyMuPDF (fitz) — fast and lightweight."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(filepath)
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return text
    except ImportError:
        # Fallback to pdfplumber if fitz not available
        try:
            import pdfplumber
            with pdfplumber.open(filepath) as pdf:
                text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            return text
        except ImportError:
            raise ImportError("Install PyMuPDF: pip install pymupdf")


def extract_text_from_docx(filepath):
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document
        doc = Document(filepath)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        return "\n".join(paragraphs)
    except ImportError:
        raise ImportError("Install python-docx: pip install python-docx")


def extract_text_from_file(filepath):
    """
    Auto-detect file type and extract text.
    Returns cleaned plain text string.
    """
    filepath_lower = filepath.lower()

    if filepath_lower.endswith('.pdf'):
        raw_text = extract_text_from_pdf(filepath)
    elif filepath_lower.endswith('.docx'):
        raw_text = extract_text_from_docx(filepath)
    else:
        raise ValueError("Unsupported file type. Use PDF or DOCX.")

    # Clean up text: normalize whitespace, remove special chars
    text = re.sub(r'\s+', ' ', raw_text)       # collapse multiple spaces
    text = re.sub(r'\n{3,}', '\n\n', raw_text) # max 2 consecutive newlines
    text = text.strip()

    return text
